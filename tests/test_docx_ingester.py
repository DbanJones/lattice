"""Tests for the DOCX outline ingester."""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from lattice.graph.models import ClaimType, RelationshipType, SectionRole
from lattice.ingester.docx import DOCXOutlineIngester, _docx_to_outline_markdown
from lattice.utils.config import Config


def _load_config(tmp_path: Path) -> Config:
    (tmp_path / "config.yml").write_text("", encoding="utf-8")
    return Config.load(tmp_path)


def _save(doc: Document, tmp_path: Path, name: str = "outline.docx") -> Path:
    path = tmp_path / name
    doc.save(str(path))
    return path


# ─── Conversion helpers ──────────────────────────────

def test_docx_converts_thesis_marker_and_sections(tmp_path: Path) -> None:
    doc = Document()
    doc.add_heading("THESIS", level=1)
    doc.add_paragraph("Forecasts diverge because of assumption, not measurement.")
    doc.add_heading("A. First section", level=1)
    doc.add_paragraph("A claim with source. [ref: koomey_2015]")
    doc.add_heading("B. Second section", level=1)
    doc.add_paragraph("MY VIEW: my synthesis claim. [user_synthesis]")

    outline = _docx_to_outline_markdown(doc)
    assert "# THESIS" in outline
    assert "Forecasts diverge" in outline
    assert "# A. First section" in outline
    assert "# B. Second section" in outline
    assert "- A claim with source" in outline
    assert "- MY VIEW: my synthesis claim" in outline


def test_docx_autonumbers_sections_without_letter_prefix(tmp_path: Path) -> None:
    doc = Document()
    doc.add_heading("THESIS", level=1)
    doc.add_paragraph("A short thesis.")
    doc.add_heading("Context", level=1)  # no letter prefix
    doc.add_paragraph("Claim in the first section.")
    doc.add_heading("Gap analysis", level=1)  # no letter prefix
    doc.add_paragraph("Claim in the second section.")
    outline = _docx_to_outline_markdown(doc)
    assert "# A. Context" in outline
    assert "# B. Gap analysis" in outline


def test_docx_preamble_paragraph_becomes_thesis(tmp_path: Path) -> None:
    doc = Document()
    # No explicit THESIS heading — the first paragraph becomes the thesis.
    doc.add_paragraph("This is implicitly the thesis.")
    doc.add_heading("A. First section", level=1)
    doc.add_paragraph("A claim.")
    outline = _docx_to_outline_markdown(doc)
    assert "# THESIS" in outline
    assert "This is implicitly the thesis" in outline


# ─── End-to-end ingester ─────────────────────────────

async def test_docx_ingester_end_to_end(tmp_path: Path) -> None:
    doc = Document()
    doc.add_heading("THESIS", level=1)
    doc.add_paragraph(
        "ICT energy forecasts diverge by twenty-fold because each embeds "
        "untested assumptions."
    )
    doc.add_heading("A. The forecast landscape", level=1)
    doc.add_paragraph(
        "Stabilisation camp assumes efficiency offsets growth [ref: masanet_2020]"
    )
    doc.add_paragraph(
        "MY VIEW: divergence is assumption-driven, not measurement-driven "
        "[user_synthesis] [supports: thesis]"
    )
    doc.add_heading("B. Gap 1: efficiency assumptions", level=1)
    doc.add_paragraph(
        "Koomey's Law doubling period lengthened [ref: koomey_2015] [role: evidence]"
    )
    doc.add_paragraph(
        "COUNTER: accelerator-era architectures partially compensate "
        "[user_synthesis] [weak]"
    )
    path = _save(doc, tmp_path)

    config = _load_config(tmp_path)
    graph = await DOCXOutlineIngester(config).ingest(path, project_name="test_docx")

    # Thesis extracted
    assert graph.thesis_statement and "twenty-fold" in graph.thesis_statement
    thesis_claim = next(c for c in graph.claims if c.claim_id == "cl.thesis")
    assert thesis_claim.type == ClaimType.user_synthesis

    # Two body sections + thesis section
    section_ids = [s.section_id for s in graph.sections]
    assert section_ids == ["s.thesis", "s.a", "s.b"]

    # MY VIEW claim has supports-thesis relationship
    my_view = next(c for c in graph.claims if "divergence is assumption" in c.statement)
    supports = [r for r in graph.relationships if r.from_claim == my_view.claim_id]
    assert any(r.type == RelationshipType.supports and r.to_claim == "cl.thesis" for r in supports)

    # COUNTER claim has contradicts-thesis relationship
    counter = next(c for c in graph.claims if "accelerator-era" in c.statement)
    contradicts = [r for r in graph.relationships if r.from_claim == counter.claim_id]
    assert any(r.type == RelationshipType.contradicts and r.to_claim == "cl.thesis" for r in contradicts)

    # [ref:] tag populated evidence
    koomey_claim = next(c for c in graph.claims if "Koomey" in c.statement)
    assert any(ev.source == "koomey_2015" for ev in koomey_claim.evidence)


async def test_docx_ingester_preserves_role_tags(tmp_path: Path) -> None:
    doc = Document()
    doc.add_heading("THESIS", level=1)
    doc.add_paragraph("Thesis.")
    doc.add_heading("A. Section", level=1)
    doc.add_paragraph("A setup claim [ref: x] [role: setup]")
    doc.add_paragraph("A conclusion claim [role: conclusion]")
    path = _save(doc, tmp_path)

    config = _load_config(tmp_path)
    graph = await DOCXOutlineIngester(config).ingest(path, project_name="t")
    setup = next(c for c in graph.claims if "setup claim" in c.statement)
    assert "role:setup" in setup.tags
    conclusion = next(c for c in graph.claims if "conclusion claim" in c.statement)
    assert "role:conclusion" in conclusion.tags
