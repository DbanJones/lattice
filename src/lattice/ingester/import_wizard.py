"""Import-an-existing-paper wizard.

Lowers the on-ramp from "30 minutes of setup" to one command. Takes
a DOCX or markdown draft, optionally a Zotero / BibTeX library, and
produces a project ready to ingest.

Detection logic:

1. **Lattice-format markdown** — already has `# THESIS` / `# A.` /
   `## A.1` headings + bulleted claims. Lands directly in
   ``structure/outline.md``; the ingester will read it as-is.
2. **Structured DOCX** — has Heading 1 / Heading 2 / bullet styles.
   Routed through the existing ``DOCXOutlineIngester`` which converts
   to the lattice tag vocabulary.
3. **Raw prose** (DOCX or markdown) — no headings or no consistent
   structure. Saved to ``structure/outline.raw.md``; the auto-outliner
   will run on the next ``Scaffold`` activity to produce a tagged
   outline.

The wizard never destroys content. If the destination outline files
already exist, the import is refused unless ``overwrite=True``.
"""

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Literal

from .auto_outliner import looks_like_lattice_outline


# Markdown heading detector — same pattern the markdown ingester uses,
# matched without a depth check so we can quickly say "this looks
# structured."
_LATTICE_HEADING_RE = re.compile(
    r"^\s*(?:#{1,3})\s+(?:THESIS|[A-Z](?:\.\d+)*\.?)\s",
    re.MULTILINE,
)


# ─── public entry point ──────────────────────────────


@dataclass
class ImportResult:
    """What the wizard did."""

    project_path: Path
    document_kind: Literal["lattice_md", "raw_md", "structured_docx", "raw_docx"]
    outline_destination: Path  # where the parsed/extracted text landed
    raw_archive: Path | None = None  # outline.raw.md when prose
    sources_imported: int = 0
    sources_duplicates: int = 0
    next_steps: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)


def import_paper(
    document: Path,
    project_path: Path,
    *,
    voice_name: str = "academic",
    references_file: Path | None = None,
    overwrite: bool = False,
) -> ImportResult:
    """Set up ``project_path`` from ``document`` (+ optional Zotero /
    BibTeX library). Returns an ``ImportResult`` describing what
    landed and what the user should do next.

    The wizard creates the project directory if needed, copies the
    canonical academic voice file into ``voices/``, decides whether
    the document needs auto-outlining, and (when given) imports a
    reference library into the source store.
    """
    document = document.resolve()
    if not document.exists():
        raise FileNotFoundError(f"Document not found: {document}")
    project_path = project_path.resolve()

    # 1. Scaffold the project tree.
    _scaffold_project(project_path, voice_name)

    # 2. Decide what kind of document this is and route accordingly.
    suffix = document.suffix.lower()
    structure_dir = project_path / "structure"
    if suffix in (".md", ".markdown", ".txt"):
        result = _import_markdown(
            document, structure_dir, project_path, overwrite=overwrite,
        )
    elif suffix == ".docx":
        result = _import_docx(
            document, structure_dir, project_path, overwrite=overwrite,
        )
    else:
        raise ValueError(
            f"Unsupported document type: {suffix}. "
            f"Provide a .md, .markdown, .txt, or .docx file."
        )

    # 3. Optional reference import.
    if references_file is not None:
        added, dupes = _import_references(project_path, references_file)
        result.sources_imported = added
        result.sources_duplicates = dupes

    # 4. Compose next-step guidance based on what landed.
    result.next_steps = _build_next_steps(result, voice_name)
    return result


# ─── project scaffolding ─────────────────────────────


def _scaffold_project(project_path: Path, voice_name: str) -> None:
    """Create the standard project layout. Idempotent — safe to call
    on an existing project; never overwrites existing files."""
    (project_path / "structure").mkdir(parents=True, exist_ok=True)
    (project_path / "refs" / "papers").mkdir(parents=True, exist_ok=True)
    (project_path / "refs" / "notes").mkdir(parents=True, exist_ok=True)
    (project_path / "refs" / "data").mkdir(parents=True, exist_ok=True)
    (project_path / "voices").mkdir(parents=True, exist_ok=True)
    (project_path / "outputs").mkdir(parents=True, exist_ok=True)
    (project_path / ".lattice").mkdir(parents=True, exist_ok=True)

    # Drop a config.yml if absent.
    config_path = project_path / "config.yml"
    if not config_path.exists():
        config_path.write_text(
            f"default_voice: {voice_name}\n"
            f"# autocorrect: safe   # safe / aggressive / none\n",
            encoding="utf-8",
        )

    # Copy the canonical voice file from the package examples.
    voice_target = project_path / "voices" / f"{voice_name}.voice.md"
    if not voice_target.exists():
        # Examples directory lives at <package_root>/../examples/voices.
        package_root = Path(__file__).resolve().parents[2]
        candidates = [
            package_root.parent / "examples" / "voices" / f"{voice_name}.voice.md",
            package_root / "examples" / "voices" / f"{voice_name}.voice.md",
        ]
        for source in candidates:
            if source.exists():
                shutil.copy2(source, voice_target)
                break


# ─── markdown handling ───────────────────────────────


def _import_markdown(
    document: Path,
    structure_dir: Path,
    project_path: Path,
    *,
    overwrite: bool,
) -> ImportResult:
    text = document.read_text(encoding="utf-8", errors="replace")
    target_md = structure_dir / "outline.md"
    target_raw = structure_dir / "outline.raw.md"

    if looks_like_lattice_outline(text):
        # Lattice-format: copy directly to outline.md.
        if target_md.exists() and not overwrite:
            raise FileExistsError(
                f"{target_md} already exists. Pass overwrite=True to replace."
            )
        target_md.write_text(text, encoding="utf-8")
        return ImportResult(
            project_path=project_path,
            document_kind="lattice_md",
            outline_destination=target_md,
        )

    # Raw prose: archive to outline.raw.md so the auto-outliner can
    # convert later. Don't touch outline.md — when auto-outliner runs,
    # it'll write the structured form there.
    if target_raw.exists() and not overwrite:
        raise FileExistsError(
            f"{target_raw} already exists. Pass overwrite=True to replace."
        )
    target_raw.write_text(text, encoding="utf-8")
    return ImportResult(
        project_path=project_path,
        document_kind="raw_md",
        outline_destination=target_raw,
        raw_archive=target_raw,
    )


# ─── DOCX handling ───────────────────────────────────


def _import_docx(
    document: Path,
    structure_dir: Path,
    project_path: Path,
    *,
    overwrite: bool,
) -> ImportResult:
    """DOCX path. We try the structured DOCX ingester first; if it
    produces meaningful headings, the doc is structured. Otherwise we
    fall back to plain-text extraction → outline.raw.md."""
    try:
        from docx import Document  # python-docx
    except ImportError as e:
        raise ImportError(
            "python-docx is required to import DOCX files. "
            "Install with `uv pip install python-docx`."
        ) from e

    doc = Document(str(document))
    structured_text = _docx_to_lattice_text(doc)
    if structured_text and _LATTICE_HEADING_RE.search(structured_text):
        # Structured: write directly to outline.md.
        target_md = structure_dir / "outline.md"
        if target_md.exists() and not overwrite:
            raise FileExistsError(
                f"{target_md} already exists. Pass overwrite=True to replace."
            )
        target_md.write_text(structured_text, encoding="utf-8")
        return ImportResult(
            project_path=project_path,
            document_kind="structured_docx",
            outline_destination=target_md,
        )

    # Raw DOCX: dump plain text → outline.raw.md.
    plain = _docx_to_plain_text(doc)
    target_raw = structure_dir / "outline.raw.md"
    if target_raw.exists() and not overwrite:
        raise FileExistsError(
            f"{target_raw} already exists. Pass overwrite=True to replace."
        )
    target_raw.write_text(plain, encoding="utf-8")
    return ImportResult(
        project_path=project_path,
        document_kind="raw_docx",
        outline_destination=target_raw,
        raw_archive=target_raw,
    )


def _docx_to_lattice_text(doc) -> str:
    """Walk a python-docx Document and emit lattice-style markdown.

    Reuses the heuristic from the existing DOCX ingester: Heading 1 /
    Heading 2 → top-level / nested section markers; everything else
    → bulleted claim line.
    """
    out: list[str] = []
    section_letter = ord("A") - 1
    sub_index: dict[int, int] = {}  # depth → counter
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style = (para.style.name if para.style else "").lower()
        if "heading 1" in style or style == "title":
            section_letter += 1
            sub_index.clear()
            letter = chr(section_letter)
            # Strip any pre-existing "A. " prefix so we don't double up.
            text = re.sub(r"^[A-Z]\.\s+", "", text)
            out.append(f"\n# {letter}. {text}\n")
        elif "heading 2" in style:
            sub_index[2] = sub_index.get(2, 0) + 1
            letter = chr(section_letter) if section_letter >= ord("A") else "A"
            out.append(f"\n## {letter}.{sub_index[2]} {text}\n")
        elif "heading 3" in style:
            sub_index[3] = sub_index.get(3, 0) + 1
            letter = chr(section_letter) if section_letter >= ord("A") else "A"
            two = sub_index.get(2, 1)
            out.append(f"\n### {letter}.{two}.{sub_index[3]} {text}\n")
        else:
            # Body / bullet: render as a claim. Auto-outliner / editor
            # can refine tags later; we don't try to guess them here.
            out.append(f"  - {text}")
    return "\n".join(out).strip() + "\n"


def _docx_to_plain_text(doc) -> str:
    """Pull plain text out of a DOCX with no structural mapping. Used
    when the DOCX has no usable headings — the auto-outliner takes
    this raw text and extracts structure on the next Scaffold pass."""
    return "\n\n".join(
        (para.text or "").strip()
        for para in doc.paragraphs
        if (para.text or "").strip()
    )


# ─── reference import ────────────────────────────────


def _import_references(
    project_path: Path, references_file: Path,
) -> tuple[int, int]:
    """Run the reference importer + merge into the project's source
    store. Returns ``(added, duplicates)``. Defensive: catches parse
    errors and returns (0, 0) rather than aborting the wizard."""
    try:
        from ..graph.store import GraphStore
        from ..references.importers import (
            import_references_from_file,
            merge_into_store,
        )
        report = import_references_from_file(references_file)
        if not report.sources:
            return 0, 0
        store = GraphStore.load(project_path)
        existing = store.list_sources()
        merged, decisions = merge_into_store(report.sources, existing)
        added_ids = {s.source_id for s in merged} - {s.source_id for s in existing}
        for src in merged:
            if src.source_id in added_ids:
                store.save_source(src)
        added = sum(1 for v in decisions.values() if v == "added")
        dupes = sum(1 for v in decisions.values() if v == "duplicate")
        return added, dupes
    except Exception:  # noqa: BLE001 - defensive; user can re-run separately
        return 0, 0


# ─── next-step messaging ─────────────────────────────


def _build_next_steps(result: ImportResult, voice_name: str) -> list[str]:
    """Produce a concrete list of "what to do now" instructions
    tailored to the document kind."""
    steps: list[str] = []
    project = result.project_path.name
    if result.document_kind in ("lattice_md", "structured_docx"):
        steps.append(f"Run `lattice ingest {project}` to parse the outline.")
        if result.sources_imported == 0:
            steps.append(
                f"Drop reference PDFs into {result.project_path}/refs/papers/, "
                f"or run `lattice references import {project} <file>` to add a "
                "Zotero / BibTeX library."
            )
        else:
            steps.append(
                f"Run `lattice index {project}` if you also have PDFs in "
                "refs/papers/."
            )
        steps.append(
            f"Run `lattice annotate {project}` (LLM) to fill in claim types, "
            "roles, and inline-citation extraction."
        )
        steps.append(
            f"Then `lattice render {project} --voice {voice_name}` for a draft."
        )
    else:
        # Raw prose path.
        steps.append(
            f"Open {result.outline_destination} and check the imported text."
        )
        steps.append(
            f"Run the Scaffold activity (web UI) or `lattice serve` and click "
            f"Scaffold → Thorough → Start to auto-outline this paper."
        )
        steps.append(
            "Once scaffolded, the standard workflow applies: ingest → "
            "annotate → render → audit."
        )
    return steps
