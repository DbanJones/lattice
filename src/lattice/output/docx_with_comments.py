"""Write the rendered paper as a DOCX with unresolved flags as native Word comments.

python-docx has no public comments API, so we:

1. Build a minimal DOCX from the rendered markdown using python-docx, with
   `commentRangeStart` / `commentRangeEnd` / `commentReference` XML
   elements embedded in the paragraphs that need comments.
2. Save the DOCX (it's an Open XML / Office Open XML zip bundle).
3. Post-process the saved zip to add `word/comments.xml`, update
   `[Content_Types].xml`, and wire the relationship in
   `word/_rels/document.xml.rels`.

Markdown parsing is intentionally shallow — we only need headings (`#`,
`##`), paragraphs, and simple pipe tables. Lattice's rendered prose stays
structurally simple, so a full CommonMark engine would be overkill.
"""

from __future__ import annotations

import re
import shutil
import zipfile
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from ..graph.models import AuditFlag


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_COMMENTS_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
_COMMENTS_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)

ET.register_namespace("w", _W_NS)


def write_paper_with_flags(
    markdown_text: str,
    unresolved_flags: list[AuditFlag],
    output_path: Path,
) -> tuple[Path, int]:
    """Build a DOCX from the rendered markdown; attach unresolved flags as Word comments.

    Returns (output_path, number_of_comments_attached).
    """
    doc = Document()
    paragraphs = _render_markdown_to_docx(doc, markdown_text)

    attacher = _CommentAttacher()
    for flag in unresolved_flags:
        attacher.attach(flag, paragraphs)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    if attacher.comment_count > 0:
        _postprocess_zip_with_comments(output_path, attacher.comments_xml_bytes())
    return output_path, attacher.comment_count


# ─── markdown → docx ───────────────────────────────

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")


def _render_markdown_to_docx(doc: Document, text: str) -> list:
    paragraphs: list = []
    buffer: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    def flush_paragraph() -> None:
        if not buffer:
            return
        combined = " ".join(line.strip() for line in buffer).strip()
        if combined:
            p = doc.add_paragraph(combined)
            paragraphs.append(p)
        buffer.clear()

    def flush_table() -> None:
        nonlocal in_table
        rows = [r for r in table_rows if any(c.strip() for c in r)]
        if rows:
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    if ci < cols:
                        table.cell(ri, ci).text = cell.strip()
        table_rows.clear()
        in_table = False

    for line in text.splitlines():
        stripped = line.rstrip()

        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_paragraph()
            in_table = True
            cells = [c for c in stripped.strip("|").split("|")]
            if all(set(c.strip()) <= {"-", ":", " "} for c in cells):
                continue
            table_rows.append(cells)
            continue
        elif in_table:
            flush_table()

        heading = _HEADING_RE.match(stripped)
        if heading:
            flush_paragraph()
            level = min(len(heading.group(1)), 4)
            p = doc.add_heading(heading.group(2).strip(), level=level)
            paragraphs.append(p)
            continue

        if not stripped:
            flush_paragraph()
            continue

        buffer.append(stripped)

    flush_paragraph()
    if in_table:
        flush_table()
    return paragraphs


# ─── comment XML building ───────────────────────────

class _CommentAttacher:
    def __init__(self) -> None:
        self._comments_root = etree.Element(qn("w:comments"), nsmap={"w": _W_NS})
        self._next_id = 0

    @property
    def comment_count(self) -> int:
        return self._next_id

    def attach(self, flag: AuditFlag, paragraphs: list) -> bool:
        target = (flag.offending_text or "").strip()
        if not target:
            return False

        for paragraph in paragraphs:
            if target in paragraph.text:
                self._wrap(paragraph, target, flag)
                return True

        # Progressive partial-match fallback. Edits often trim a tail off
        # the flagged span; try successively shorter prefixes until one
        # matches, down to a floor of 15 chars.
        for trim in (60, 40, 25, 15):
            if len(target) < trim:
                continue
            prefix = target[:trim]
            for paragraph in paragraphs:
                if prefix in paragraph.text:
                    self._wrap(paragraph, prefix, flag)
                    return True
        return False

    def _wrap(self, paragraph, needle: str, flag: AuditFlag) -> None:
        cid = self._next_id
        self._next_id += 1

        # Append the comment body to the comments XML doc.
        comment = etree.SubElement(
            self._comments_root,
            qn("w:comment"),
            {
                qn("w:id"): str(cid),
                qn("w:author"): "Lattice",
                qn("w:initials"): "LT",
                qn("w:date"): datetime.now(timezone.utc).isoformat(),
            },
        )
        lines = [
            f"[{flag.category.value} / {flag.severity.value}] {flag.rule_id}",
            flag.rule_description,
        ]
        if flag.suggestion:
            lines.append(f"Suggestion: {flag.suggestion}")
        for line in lines:
            cp = etree.SubElement(comment, qn("w:p"))
            cr = etree.SubElement(cp, qn("w:r"))
            ct = etree.SubElement(cr, qn("w:t"))
            ct.text = line
            ct.set(qn("xml:space"), "preserve")

        # Rewrite the paragraph: pre | commentRangeStart | mid | commentRangeEnd | commentReference | post
        full = paragraph.text
        idx = full.find(needle)
        if idx < 0:
            return
        pre = full[:idx]
        mid = full[idx : idx + len(needle)]
        post = full[idx + len(needle) :]

        p_elem = paragraph._element
        for r in list(p_elem.findall(qn("w:r"))):
            p_elem.remove(r)

        if pre:
            _append_run(p_elem, pre)
        etree.SubElement(p_elem, qn("w:commentRangeStart"), {qn("w:id"): str(cid)})
        _append_run(p_elem, mid)
        etree.SubElement(p_elem, qn("w:commentRangeEnd"), {qn("w:id"): str(cid)})

        ref_r = etree.SubElement(p_elem, qn("w:r"))
        rpr = etree.SubElement(ref_r, qn("w:rPr"))
        etree.SubElement(rpr, qn("w:rStyle"), {qn("w:val"): "CommentReference"})
        etree.SubElement(ref_r, qn("w:commentReference"), {qn("w:id"): str(cid)})

        if post:
            _append_run(p_elem, post)

    def comments_xml_bytes(self) -> bytes:
        return etree.tostring(
            self._comments_root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )


def _append_run(paragraph_elem, text: str) -> None:
    r = etree.SubElement(paragraph_elem, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.set(qn("xml:space"), "preserve")
    t.text = text


# ─── post-process the saved .docx zip ──────────────

def _postprocess_zip_with_comments(docx_path: Path, comments_xml: bytes) -> None:
    """Add comments.xml + content-type entry + rel entry to the saved DOCX."""
    # Read all existing parts into memory.
    with zipfile.ZipFile(docx_path, "r") as zf:
        parts = {name: zf.read(name) for name in zf.namelist()}

    # Add the comments part.
    parts["word/comments.xml"] = comments_xml

    # Update [Content_Types].xml to register the comments content type.
    ct_name = "[Content_Types].xml"
    ct_xml = parts[ct_name]
    ct_root = etree.fromstring(ct_xml)
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    # Check for an existing Override for /word/comments.xml.
    found = any(
        child.get("PartName") == "/word/comments.xml"
        for child in ct_root.findall(f"{{{ct_ns}}}Override")
    )
    if not found:
        override = etree.SubElement(
            ct_root,
            f"{{{ct_ns}}}Override",
            PartName="/word/comments.xml",
            ContentType=_COMMENTS_CONTENT_TYPE,
        )
    parts[ct_name] = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    # Wire the relationship from document.xml to comments.xml.
    rels_name = "word/_rels/document.xml.rels"
    rels_xml = parts[rels_name]
    rels_root = etree.fromstring(rels_xml)
    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    existing_rids = {
        child.get("Id")
        for child in rels_root.findall(f"{{{rels_ns}}}Relationship")
    }
    # Find a unique rId not already used.
    n = 1
    while f"rId{n}" in existing_rids:
        n += 1
    etree.SubElement(
        rels_root,
        f"{{{rels_ns}}}Relationship",
        Id=f"rId{n}",
        Type=_COMMENTS_REL_TYPE,
        Target="comments.xml",
    )
    parts[rels_name] = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    # Rewrite the zip with the updated set of parts.
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in parts.items():
            zf.writestr(name, data)
    shutil.move(str(tmp_path), str(docx_path))
