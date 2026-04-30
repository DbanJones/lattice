"""DOCX indexer using python-docx.

Passage IDs: p.<paragraph_seq> — sequential within the document.
Preserves heading level in PassageLocation.section as "Hn" and retains
the outline order of paragraphs.
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document

from ..graph.models import (
    Citation,
    Passage,
    PassageLocation,
    PassageType,
    Source,
    SourceMetadata,
    SourceType,
)
from .base import Indexer


class DOCXIndexer(Indexer):
    def index(self, file_path: Path) -> Source:
        doc = Document(str(file_path))
        passages: list[Passage] = []
        current_heading = ""

        for seq, para in enumerate(doc.paragraphs, start=1):
            text = (para.text or "").strip()
            if not text:
                continue
            style_name = (para.style.name if para.style else "") or ""
            is_heading = style_name.lower().startswith("heading")
            if is_heading:
                current_heading = text

            passages.append(
                Passage(
                    id=f"p.{seq}.1",
                    text=text,
                    location=PassageLocation(paragraph=seq, section=current_heading or None),
                    type=PassageType.claim,
                    char_count=len(text),
                )
            )

        title = _infer_title(doc, file_path)

        return Source(
            source_id=Indexer.slugify(file_path.stem),
            type=SourceType.primary_paper,
            citation=Citation(authors=[], year=None, title=title),
            passages=passages,
            metadata=SourceMetadata(
                date_added=datetime.now(timezone.utc),
                file_path=str(file_path),
                hash="",
            ),
        )


def _infer_title(doc, file_path: Path) -> str:
    # First paragraph with a "Title" or "Heading 1" style wins.
    for para in doc.paragraphs:
        style = (para.style.name if para.style else "").lower()
        if style in {"title", "heading 1"} and (para.text or "").strip():
            return para.text.strip()
    return file_path.stem
